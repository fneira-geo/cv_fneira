#!/usr/bin/env python3
"""Genera template_ciren_blank.docx — réplica plana del formato CIREN."""
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT     = Path(__file__).parent.resolve()
OUT_PATH = ROOT / "template_ciren" / "template_ciren_blank.docx"


def set_single_spacing(paragraph):
    """Set line spacing to exactly 1.0 (single)."""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def setup_document() -> Document:
    """Create document with standard margins and default font."""
    doc = Document()
    for s in doc.sections:
        s.top_margin    = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin   = Cm(3)
        s.right_margin  = Cm(3)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return doc


def add_heading(doc, text: str) -> None:
    """Add section heading: bold + underline."""
    p = doc.add_paragraph()
    set_single_spacing(p)
    r = p.add_run(text)
    r.bold = True
    r.underline = True
    r.font.name = "Calibri"
    r.font.size = Pt(11)


def add_subheading(doc, text: str) -> None:
    """Add subheading: normal text (no bold, no underline)."""
    p = doc.add_paragraph(text)
    set_single_spacing(p)


def set_table_borders_black(table) -> None:
    """Add simple black borders to table."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tbl_pr.append(borders)


def add_section_break(doc) -> None:
    """Add empty paragraph for spacing between sections."""
    p = doc.add_paragraph()
    set_single_spacing(p)


def add_invisible_table(doc, filas: list[tuple[str, str]]):
    """2-column table with invisible borders: Label | Value."""
    t = doc.add_table(rows=len(filas), cols=2)
    tbl_pr = t._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "none")
        borders.append(b)
    tbl_pr.append(borders)
    for i, (label, valor) in enumerate(filas):
        row = t.rows[i]
        r0 = row.cells[0].paragraphs[0].add_run(label)
        r0.bold = True
        r0.font.name = "Calibri"
        r0.font.size = Pt(11)
        r1 = row.cells[1].paragraphs[0].add_run(valor)
        r1.font.name = "Calibri"
        r1.font.size = Pt(11)
    return t


def add_single_cell_table(doc):
    """1-column table (single cell) for specialization section."""
    t = doc.add_table(rows=1, cols=1)
    set_table_borders_black(t)
    return t


def add_empty_table(doc, headers: list, n_filas: int = 3):
    """Add table with headers and empty rows."""
    t = doc.add_table(rows=1 + n_filas, cols=len(headers))
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    set_table_borders_black(t)
    return t


def add_label(doc, label: str) -> None:
    """Add bold label with colon."""
    p = doc.add_paragraph()
    r = p.add_run(f"{label}:")
    r.bold = True


def add_bullet(doc, text: str = "", level: int = 0) -> None:
    """Add bullet point (empty or with text)."""
    style = "List Bullet" if level == 0 else "List Bullet 2"
    doc.add_paragraph(text, style=style)


# ── secciones ───────────────────────────────────────────────────────

def section_datos_contacto(doc):
    add_heading(doc, "Datos de contacto")
    filas = [
        ("Nombre completo", ""),
        ("Fecha de nacimiento", ""),
        ("RUT", ""),
        ("Estado civil", ""),
        ("Nacionalidad", ""),
        ("Dirección", ""),
        ("Teléfono", ""),
        ("Correo electrónico", ""),
    ]
    add_invisible_table(doc, filas)


def section_antecedentes(doc):
    add_heading(doc, "Antecedentes Académicos")
    add_empty_table(doc, ["Título/Grado/Especialización", "Institución", "Año de Obtención"])


def section_linea_especializacion(doc):
    add_heading(doc, "Línea de especialización")
    add_single_cell_table(doc)


def section_experiencia(doc):
    add_heading(doc, "Experiencia laboral últimos 5 años")
    add_empty_table(doc, ["Nombre del cargo", "Funciones del cargo",
                          "Organismo/Institución", "Período"])


def section_proyectos(doc):
    add_heading(doc, "Participación en proyectos últimos 5 años")
    add_empty_table(doc, ["Cargo", "Organismo/Institución", "Período"])


def section_competencias(doc):
    add_heading(doc, "Competencias profesionales y manejo de software")

    add_subheading(doc, "Capacitaciones")
    add_bullet(doc, "", level=0)
    add_bullet(doc, "", level=1)
    add_bullet(doc, "", level=1)

    add_subheading(doc, "Programación")
    add_bullet(doc, "")
    add_bullet(doc, "")

    add_subheading(doc, "Software GIS")
    add_bullet(doc, "")
    add_bullet(doc, "")

    add_subheading(doc, "Idiomas")
    add_bullet(doc, "")
    add_bullet(doc, "")

    add_subheading(doc, "Otros")
    add_bullet(doc, "")
    add_bullet(doc, "")


def main() -> None:
    doc = setup_document()

    # Title: blank centered bold (user fills in their name)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("").bold = True

    add_section_break(doc)
    section_datos_contacto(doc)
    add_section_break(doc)
    section_antecedentes(doc)
    add_section_break(doc)
    section_linea_especializacion(doc)
    add_section_break(doc)
    section_experiencia(doc)
    add_section_break(doc)
    section_proyectos(doc)
    add_section_break(doc)
    section_competencias(doc)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"  *  {OUT_PATH.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
