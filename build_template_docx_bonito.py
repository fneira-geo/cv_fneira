#!/usr/bin/env python3
"""Genera template_ciren_blank.docx imitando el formato CIREN."""
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT       = Path(__file__).parent.resolve()
OUT_PATH   = ROOT / "template_ciren" / "template_ciren_blank.docx"

CIREN_TEAL = RGBColor(0x0F, 0x47, 0x61)
GRAY       = RGBColor(0x80, 0x80, 0x80)
FONT       = "Calibri"


# ── helpers de estilo ───────────────────────────────────────────────

def setup_document() -> Document:
    """Set up document with margins and default font."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(12)
    return doc


def set_lang_es_cl(doc):
    """Set document language to Spanish (Chile)."""
    styles_element = doc.styles.element
    docDefaults = styles_element.find(qn('w:docDefaults'))
    if docDefaults is None:
        docDefaults = OxmlElement('w:docDefaults')
        styles_element.insert(0, docDefaults)

    rPr = docDefaults.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        docDefaults.append(rPr)

    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        rPr.append(lang)
    lang.set(qn('w:val'), 'es-CL')


def add_title(doc, text):
    """Add title: 20pt bold teal centred."""
    p = doc.add_paragraph(text)
    p.alignment = 1  # Center
    for run in p.runs:
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = CIREN_TEAL


def add_heading(doc, text):
    """Add heading: 14pt bold teal with bottom border."""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = CIREN_TEAL

    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '0F4761')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_subheading(doc, text):
    """Add subheading: 12pt bold."""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.bold = True


def add_placeholder_run(p, txt):
    """Add italic gray placeholder text [...]."""
    run = p.add_run(txt)
    run.font.italic = True
    run.font.color.rgb = GRAY


def add_label_value(doc, label, hint):
    """Add label-value pair: bold label + placeholder hint."""
    p = doc.add_paragraph()

    label_run = p.add_run(f"{label}: ")
    label_run.font.bold = True
    label_run.font.color.rgb = RGBColor(0, 0, 0)

    add_placeholder_run(p, hint)


def add_table(doc, headers, n_filas, hint_por_columna):
    """Add table with headers, n rows, and column hints."""
    table = doc.add_table(rows=n_filas + 1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        header_cells[i].text = header_text
        # Bold white text
        for para in header_cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

        # Teal background
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '0F4761')
        header_cells[i]._element.get_or_add_tcPr().append(shd)

    # Data rows with hints
    for row_idx in range(1, n_filas + 1):
        row = table.rows[row_idx]
        for col_idx, hint in enumerate(hint_por_columna):
            cell = row.cells[col_idx]
            p = cell.paragraphs[0]
            add_placeholder_run(p, hint)

    # Add table borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def add_bullet(doc, text, level=0):
    """Add bullet point with optional nesting (level 0 or 1)."""
    p = doc.add_paragraph(text)
    if level == 0:
        p.style = 'List Bullet'
    elif level == 1:
        p.style = 'List Bullet 2'

    # Make text italic gray placeholder
    for run in p.runs:
        run.font.italic = True
        run.font.color.rgb = GRAY


# ── secciones ───────────────────────────────────────────────────────

def section_datos_contacto(doc):
    add_heading(doc, "Datos de contacto")
    for label, hint in [
        ("Nombre completo",       "[Nombre y apellidos]"),
        ("Fecha de nacimiento",   "[DD/MM/AAAA]"),
        ("RUT",                   "[XX.XXX.XXX-X]"),
        ("Estado civil",          "[Soltero/Casado/...]"),
        ("Nacionalidad",          "[País]"),
        ("Dirección",             "[Calle, número, comuna, ciudad]"),
        ("Teléfono",              "[+56 9 XXXX XXXX]"),
        ("Correo electrónico",    "[correo@dominio.cl]"),
    ]:
        add_label_value(doc, label, hint)


def section_antecedentes(doc):
    add_heading(doc, "Antecedentes Académicos")
    add_table(
        doc,
        headers=["Título / Grado / Especialización", "Institución", "Año de Obtención"],
        n_filas=3,
        hint_por_columna=["[Título obtenido]", "[Universidad / centro]", "[AAAA]"],
    )


def section_linea_especializacion(doc):
    add_heading(doc, "Línea de especialización")
    p = doc.add_paragraph()
    add_placeholder_run(p,
        "[Párrafo de máximo 5 líneas describiendo el área de especialización profesional. "
        "Resaltar en negrita las ")
    r = p.add_run("palabras clave")
    r.bold = True
    add_placeholder_run(p,
        " que identifican la experticia: hidrología, modelación, SIG, teledetección, etc.]")


def section_experiencia(doc):
    add_heading(doc, "Experiencia laboral últimos 5 años")
    add_table(
        doc,
        headers=["Nombre del cargo", "Funciones del cargo",
                 "Organismo / Institución", "Período"],
        n_filas=3,
        hint_por_columna=["[Cargo]", "[Funciones principales]",
                          "[Empresa / institución]", "[MM/AAAA – MM/AAAA]"],
    )


def section_proyectos(doc):
    add_heading(doc, "Participación en proyectos últimos 5 años")
    add_table(
        doc,
        headers=["Cargo", "Organismo / Institución", "Período"],
        n_filas=3,
        hint_por_columna=["[Rol en el proyecto]",
                          "[Empresa / institución mandante]",
                          "[MM/AAAA – MM/AAAA]"],
    )


def section_competencias(doc):
    add_heading(doc, "Competencias profesionales y manejo de software")

    add_subheading(doc, "Capacitaciones")
    add_bullet(doc, "[Nombre del curso o capacitación]", level=0)
    add_bullet(doc, "[Institución – año]",                level=1)
    add_bullet(doc, "[Duración / certificación]",         level=1)
    add_bullet(doc, "[Otra capacitación relevante]",      level=0)

    add_subheading(doc, "Programación")
    add_bullet(doc, "[Lenguaje 1 — nivel]")
    add_bullet(doc, "[Lenguaje 2 — nivel]")

    add_subheading(doc, "Software GIS")
    add_bullet(doc, "[QGIS / ArcGIS / ENVI / etc. — nivel]")

    add_subheading(doc, "Idiomas")
    add_bullet(doc, "[Idioma — nivel hablado / escrito]")

    add_subheading(doc, "Otros")
    add_bullet(doc, "[Otras herramientas, software o competencias relevantes]")


def main() -> None:
    doc = setup_document()
    set_lang_es_cl(doc)
    add_title(doc, "NOMBRE APELLIDOS")
    p = doc.add_paragraph()
    add_placeholder_run(p, "[Profesión / título profesional]")
    section_datos_contacto(doc)
    section_antecedentes(doc)
    section_linea_especializacion(doc)
    section_experiencia(doc)
    section_proyectos(doc)
    section_competencias(doc)
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"  *  {OUT_PATH.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
