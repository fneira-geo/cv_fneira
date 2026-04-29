#!/usr/bin/env python3
"""Genera CV_FernandoNeira_CIREN.docx con datos reales desde data/*.yaml."""
import re
from datetime import date
from pathlib import Path
import yaml
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from build_template_docx import (
    setup_document, add_heading, add_subheading, add_section_break,
    add_empty_table, add_bullet, add_invisible_table, add_single_cell_table,
    set_single_spacing,
)

ROOT      = Path(__file__).parent.resolve()
DATA_DIR  = ROOT / "data"
OUT_PATH  = ROOT / "output" / "CV_FernandoNeira_CIREN.docx"

VENTANA_ANIOS = 5
HOY           = date.today()
ANIO_CORTE    = HOY.year - VENTANA_ANIOS    # 2021 cuando HOY=2026

PALABRAS_CLAVE = [
    "agroclimatología", "análisis geoespacial", "recursos hídricos",
    "geomática", "modelación de demanda hídrica", "teledetección",
    "Ciencia de Datos", "cambio climático",
]


# ── carga de datos ──────────────────────────────────────────────────
def load_yaml(name: str):
    with open(DATA_DIR / name, encoding="utf-8") as f:
        return yaml.safe_load(f)


# ── filtros ─────────────────────────────────────────────────────────
def fin_dentro_ventana(fin: str) -> bool:
    if fin == "presente":
        return True
    return fin >= f"{ANIO_CORTE:04d}-01"

def ultimo_anio(texto: str) -> int | None:
    anios = [int(y) for y in re.findall(r"\d{4}", texto)]
    return max(anios) if anios else None


# ── helpers de contenido ────────────────────────────────────────────
def add_paragraph_with_bold(doc, text: str, keywords: list[str]):
    """Add paragraph; runs matching any keyword (case-insensitive) are bold."""
    p = doc.add_paragraph()
    set_single_spacing(p)
    pattern = re.compile("(" + "|".join(re.escape(k) for k in keywords) + ")", re.IGNORECASE)
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = p.add_run(text[pos:m.start()])
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        run = p.add_run(text[m.start():m.end()])
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        pos = m.end()
    if pos < len(text):
        run = p.add_run(text[pos:])
        run.font.name = "Calibri"
        run.font.size = Pt(11)

def fill_table_row(table, row_idx: int, valores: list[str]):
    for i, v in enumerate(valores):
        cell = table.rows[row_idx].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(v)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


# ── secciones ───────────────────────────────────────────────────────
def section_datos_contacto(doc, perfil):
    add_heading(doc, "Datos de contacto")
    filas = [
        ("Nombre completo", perfil.get("nombre", "")),
        ("Fecha de nacimiento", ""),
        ("RUT", ""),
        ("Estado civil", ""),
        ("Nacionalidad", ""),
        ("Dirección", perfil.get("ubicacion", "")),
        ("Teléfono", perfil.get("telefono", "")),
        ("Correo electrónico", perfil.get("email", "")),
    ]
    add_invisible_table(doc, filas)

def section_antecedentes(doc, educacion):
    add_heading(doc, "Antecedentes Académicos")
    headers = ["Título/Grado/Especialización", "Institución", "Año de Obtención"]
    t = add_empty_table(doc, headers, n_filas=len(educacion))
    for i, edu in enumerate(educacion, start=1):
        titulo = edu["titulo"]
        if edu.get("subtitulo"):
            titulo += f" — {edu['subtitulo']}"
        anio = edu["periodo"].split("–")[-1].strip()
        fill_table_row(t, i, [titulo, edu["institucion"], anio])

def section_linea_especializacion(doc, perfil):
    add_heading(doc, "Línea de especialización")
    t = add_single_cell_table(doc)
    cell = t.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    set_single_spacing(p)
    pattern = re.compile("(" + "|".join(re.escape(k) for k in PALABRAS_CLAVE) + ")", re.IGNORECASE)
    text = perfil["resumen"].strip()
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = p.add_run(text[pos:m.start()])
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        run = p.add_run(text[m.start():m.end()])
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        pos = m.end()
    if pos < len(text):
        run = p.add_run(text[pos:])
        run.font.name = "Calibri"
        run.font.size = Pt(11)

def section_experiencia(doc, experiencia):
    add_heading(doc, "Experiencia laboral últimos 5 años")
    headers = ["Nombre del cargo", "Funciones del cargo",
               "Organismo/Institución", "Período"]
    filtrados = [e for e in experiencia if fin_dentro_ventana(str(e.get("fin", "")))]
    t = add_empty_table(doc, headers, n_filas=len(filtrados))
    for i, e in enumerate(filtrados, start=1):
        funciones = "; ".join(l.strip() for l in e.get("logros", []))
        periodo   = f"{e['inicio']} – {e['fin']}"
        fill_table_row(t, i, [e["cargo"], funciones, e["institucion"], periodo])

def section_proyectos(doc, experiencia):
    add_heading(doc, "Participación en proyectos últimos 5 años")
    headers = ["Cargo", "Organismo/Institución", "Período"]
    filas = []
    for e in experiencia:
        if not fin_dentro_ventana(str(e.get("fin", ""))):
            continue
        for proy in e.get("proyectos", []):
            ya = ultimo_anio(proy["periodo"])
            if ya is not None and ya >= ANIO_CORTE:
                filas.append((e["cargo"], e["institucion"],
                              f"{proy['nombre']} ({proy['periodo']})"))
    t = add_empty_table(doc, headers, n_filas=len(filas))
    for i, (cargo, inst, periodo) in enumerate(filas, start=1):
        fill_table_row(t, i, [cargo, inst, periodo])

def section_competencias(doc, cursos, competencias, idiomas):
    add_heading(doc, "Competencias profesionales y manejo de software")

    # Capacitaciones (filtro 5 años)
    add_subheading(doc, "Capacitaciones")
    for c in cursos:
        ya = ultimo_anio(str(c.get("fecha", "")))
        if ya is not None and ya >= ANIO_CORTE:
            add_bullet(doc, f"{c['nombre']} — {c['institucion']} ({c['fecha']})")
            if c.get("detalle"):
                add_bullet(doc, c["detalle"].strip(), level=1)

    # Programación
    add_subheading(doc, "Programación")
    grupo_prog = next(g for g in competencias if g["grupo"] == "Programación & Datos")
    for h in grupo_prog["habilidades"]:
        add_bullet(doc, h)

    # Software GIS (extraer del grupo Teledetección lo relacionado a SIG/raster)
    add_subheading(doc, "Software GIS")
    grupo_td = next(g for g in competencias
                    if g["grupo"] == "Teledetección & Observación de la Tierra")
    for h in grupo_td["habilidades"]:
        if any(k in h for k in ("SIG", "QGIS", "ArcGIS", "GDAL", "Earth Engine",
                                "Orfeo", "Whitebox", "raster")):
            add_bullet(doc, h)

    # Idiomas
    add_subheading(doc, "Idiomas")
    for idi in idiomas:
        add_bullet(doc, f"{idi['idioma']} — {idi['nivel']}")

    # Otros (resto del grupo Teledetección + Ciencia de Datos + Análisis Climático
    # como bullets sintéticos seleccionados)
    add_subheading(doc, "Otros")
    grupo_clima = next(g for g in competencias
                       if g["grupo"] == "Análisis Climático & Recursos Hídricos")
    for h in grupo_clima["habilidades"][:5]:                # limitar volumen
        add_bullet(doc, h)
    grupo_cd = next(g for g in competencias
                    if g["grupo"] == "Ciencia de Datos & Estadística")
    for h in grupo_cd["habilidades"][:3]:
        add_bullet(doc, h)


def main() -> None:
    perfil       = load_yaml("perfil.yaml")
    educacion    = load_yaml("educacion.yaml")
    experiencia  = load_yaml("experiencia.yaml")
    cursos       = load_yaml("cursos.yaml")
    idiomas      = load_yaml("idiomas.yaml")
    competencias = load_yaml("competencias.yaml")

    doc = setup_document()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(perfil["nombre"])
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(perfil["titulo"])
    r2.font.name = "Calibri"
    r2.font.size = Pt(11)

    add_section_break(doc)
    section_datos_contacto(doc, perfil)
    add_section_break(doc)
    section_antecedentes(doc, educacion)
    add_section_break(doc)
    section_linea_especializacion(doc, perfil)
    add_section_break(doc)
    section_experiencia(doc, experiencia)
    add_section_break(doc)
    section_proyectos(doc, experiencia)
    add_section_break(doc)
    section_competencias(doc, cursos, competencias, idiomas)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"  *  {OUT_PATH.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
